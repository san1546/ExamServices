# -*- coding: utf-8 -*-
import subprocess

from docx import Document
import xlrd
import xlsxwriter
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm,Pt,RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import base64
from PIL import Image
# import win32com.client
# import pythoncom
from bin.repository import *
# class Word_2_PDF(object):
#
#     def __init__(self, filepath, Debug=False):  # param Debug: 控制过程是否可视化
#         pythoncom.CoInitialize()
#         self.wordApp = win32com.client.Dispatch('word.Application')
#         self.wordApp.Visible = Debug
#         self.myDoc = self.wordApp.Documents.Open(os.path.abspath('.\\' + filepath))
#
#     def export_pdf(self, output_file_path):  # 将Word文档转化为PDF文件
#         self.myDoc.ExportAsFixedFormat(os.path.abspath('.\\' + output_file_path), 17, Item=7, CreateBookmarks=0)

def openWord(testno, chinese_name, english_name, idno, examsite, examno, seatno, subject, examtime, photo, business_id, business_type, created_by):
    document = Document('templates/准考证模板.docx')  #打开文件demo.docx
    # 查看文本框
    children = document.element.body.iter()
    child_iters = []
    tags = []
    i = 0
    for child in children:
        # print("子标记:",child.tag)
        # 通过类型判断目录
        if child.tag.endswith(('AlternateContent', 'textbox')):
            for ci in child.iter():
                # tags.append(ci.tag)
                if ci.tag.endswith(('main}r', 'main}pPr')):
                    if ci.text == '准考证号：':
                        ci.text = '准考证号：' + testno
                    if ci.text == '中文姓名：':
                        ci.text = '中文姓名：' + chinese_name
                    if ci.text == '英文姓名：':
                        ci.text = '英文姓名：' + english_name
                    if ci.text == '证件号码：':
                        ci.text = '证件号码：' + idno
                    if ci.text == '考点名称：':
                        ci.text = '考点名称：' + examsite
                    if ci.text == '考场号：':
                        ci.text = '考场号：' + examno
                    if ci.text == '座位号：':
                        ci.text = '座位号：' + seatno
                    if ci.text == 'a':
                        ci.text = subject
                    if ci.text == 'b':
                        ci.text = examtime


                    # if ci.text == 'A':

                    if ci.text is not None:
                        ci.text.replace(" ", "")
                        # print("true")
                    # print("ci.text:", ci.text)
                    child_iters.append(ci)

    text = ['']
    for ci in child_iters:
        if ci.tag.endswith('main}pPr'):
            text.append('')
        else:
            text[-1] += ci.text
        ci.text = ''
    # trans_text = ['***' + t + '***' for t in text]
    trans_text = [ t for t in text]
    # print(trans_text)
    i, k = 0, 0
    for ci in child_iters:
        if ci.tag.endswith('main}pPr'):
            i += 1
            k = 0
        elif k == 0:
            ci.text = trans_text[i]
            k = 1

    # print("table数:", len(document.tables))
    mark = str(uuid.uuid1())
    if photo:
        run = document.tables[0].cell(0, 0).paragraphs[0].add_run()
        filetype = photo[11:15].split(";")[0]
        print("filetype:", filetype)
        imgdata = base64.b64decode(photo.split(",")[1])
        # print("imgdata:", imgdata)
        if not os.path.exists('testno/' + mark):
            os.makedirs('testno/' + mark)
        file = open('testno/' + mark + '/temp.' + filetype, 'wb')
        file.write(imgdata)
        file.close()
        f = Image.open('testno/' + mark + '/temp.' + filetype)  # 你的图片文件
        f.save('testno/' + mark + '/temp.' + filetype)  # 替换掉你的图片文件
        f.close()
        run.add_picture('testno/' + mark + '/temp.' + filetype, height=Inches(1.287), width=Inches(0.9))


    tb = document.tables[1]
    # 获取表格的行
    tb_rows = tb.rows
    #读取每一行内容
    for i in range(0, len(tb_rows)):
        row_data = []
        row_cells = tb_rows[i].cells
        # 读取每一行单元格内容
        for j in range(0, len(row_cells)):
            # print(row_cells[j].text)
            # 单元格内容
            if row_cells[j].text == 'a':
                row_cells[j].text = row_cells[j].text.replace("a", str(subject))
                row_cells[j].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            elif row_cells[j].text == 'b':
                row_cells[j].text = row_cells[j].text.replace("b", str(examtime))
                row_cells[j].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            row_data.append(row_cells[j].text)
        # print(row_data)


    # rootpath = 'testno/' + examsite + '_' + mark + '/'
    rootpath = 'testno/' + mark + '/'
    if not os.path.exists(rootpath):
        os.makedirs(rootpath)
    document.save(rootpath + '/' + testno + '.docx')  # 保存文档

    docfilelist = [testno + '.docx']
    # win doc转pdf
    # for eachdocname in docfilelist:
    #     # print('路径：', os.path.join(rootpath, eachdocname))
    #     w2p = Word_2_PDF(os.path.join(rootpath, eachdocname), False)
    #     eachpdfname = eachdocname[:eachdocname.rfind('.')] + '.pdf'
    #     # print('另一个pdf路径：', os.path.join(rootpath, eachpdfname))
    #     w2p.export_pdf(os.path.join(rootpath, eachpdfname))
    #     # print("文件名：", eachpdfname)
    #     w2p.myDoc.Close()
    #     fileinfo = os.stat(os.path.join(rootpath, eachpdfname))  # 获取文件的基本信息
    #     repository = Repository()
    #     filepath_db = os.path.join(rootpath, eachpdfname).replace("\\", "/")
    #     repository.saveExamineeCardAtt(eachpdfname, filepath_db, fileinfo.st_size, 'pdf', business_id, business_type, created_by)
    # linux doc转pdf
    for eachdocname in docfilelist:
        print("eachdocname:", eachdocname)
        doc = os.path.join(rootpath, eachdocname)
        print("doc:", doc)
        cmd = 'soffice --headless --convert-to pdf'.split() + [doc] + ' --outdir '.split() + [rootpath]
        print("cmd:", cmd)
        p = subprocess.Popen(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE)
        p.wait(timeout=10)
        stdout, stderr = p.communicate()
        # eachpdfname = eachdocname[:eachdocname.rfind('.')] + '.pdf'
        if stderr:
            raise subprocess.SubprocessError(stderr)
        # fileinfo = os.stat(os.path.join(rootpath, eachpdfname))  # 获取文件的基本信息
        # repository = Repository()
        # filepath_db = os.path.join(rootpath, eachpdfname).replace("\\", "/")
        # repository.saveExamineeCardAtt(eachpdfname, filepath_db, fileinfo.st_size, 'pdf', business_id, business_type, created_by)




